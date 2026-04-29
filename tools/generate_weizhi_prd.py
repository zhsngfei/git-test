from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE = "未至 MVP PRD"
OUTPUT_PATH = Path(r"D:\codex.files\git-test\未至-MVP-PRD.docx")


SECTIONS = [
    (
        "1. 产品定位",
        [
            "未至是一款旅行前的文化触点推荐应用。",
            "用户输入目的地城市和阅读/观影偏好后，获得与该城市相关的书籍、电影或剧集推荐，以及作品中涉及的景点触点。",
            "页面目标不是生成长篇介绍，而是通过作品卡片、封面/剧照、景点图片和简短理由，让用户快速建立对这座城市的感知。",
        ],
    ),
    (
        "2. 目标用户",
        [
            "准备出行、愿意在旅行前做一点内容补课的用户。",
            "对书籍、电影、剧集、城市气质和故事感有兴趣的文艺旅行者。",
        ],
    ),
    (
        "3. 核心流程",
        [
            "用户进入首页，输入目的地城市并选择阅读/观影偏好。",
            "系统返回该城市的作品推荐结果页，按作品卡片展示。",
            "用户浏览作品、展开相关景点、查看图片和说明，并可收藏或标记想读/想看。",
        ],
    ),
    (
        "4. 页面说明",
        [
            "页面 1：首页 / 搜索页",
            "需要展示：应用名称、城市输入框、偏好选择区、搜索按钮、示例城市入口。",
            "偏好可先使用标签形式，例如：文学、电影、剧集、经典、当代、治愈、人文。",
            "页面 2：城市结果页",
            "需要展示：城市标题、当前偏好标签、6-10 个作品卡片。",
            "每个作品卡片需包含：作品名、类型、封面/海报/剧照、一句话简介、与城市的关系、推荐理由、相关景点列表、收藏按钮、想读/想看按钮。",
            "每个作品可展开最多 10 个景点。",
            "每个景点需包含：景点名、景点图片或暂无图片状态、景点简述、它在作品中的意义、地图入口。",
            "页面 3：收藏页",
            "需要展示：已收藏作品、已标记想读/想看的作品。",
        ],
    ),
    (
        "5. 内容展示规则",
        [
            "推荐结果以作品为主，不按景点做主视图。",
            "文案保持短句，不写大段城市介绍。",
            "每个作品必须清楚表达两件事：它和这座城市有什么关系；为什么值得在去之前看/读。",
            "书籍必须有封面。",
            "影视优先展示代表性剧照；如果没有合适剧照，可使用海报或封面。",
            "景点优先展示图片；如果没有图片，仍然要展示景点内容，并明确标注“暂时没有图”。",
        ],
    ),
    (
        "6. 交互与状态",
        [
            "支持单城市搜索。",
            "支持作品收藏。",
            "支持标记想读/想看。",
            "支持展开和收起景点列表。",
            "景点缺图时显示占位状态文案，不隐藏该景点。",
            "如果城市可用内容较少，允许少于 6 个作品，并提示“当前这座城市可用的文化触点还比较少”。",
        ],
    ),
    (
        "7. 视觉方向",
        [
            "整体气质偏安静、克制、有文学感，不做常规旅游攻略风格。",
            "视觉重心放在封面、剧照、景点图上，减少大段文字堆叠。",
            "卡片应有明显层级，让用户先被图像吸引，再快速读完理由和触点。",
            "结果页应体现浏览感和收藏感，而不是工具后台感。",
        ],
    ),
    (
        "8. 交付边界",
        [
            "本次只需要输出 UI 页面设计，不需要实现前端逻辑或后端接口。",
            "不需要设计行程规划、打卡回忆册、社区发帖、多城市路线等扩展功能。",
        ],
    ),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section):
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    if "Body Accent" not in styles:
        body_style = styles.add_style("Body Accent", WD_STYLE_TYPE.PARAGRAPH)
    else:
        body_style = styles["Body Accent"]
    body_style.base_style = styles["Normal"]
    body_style.font.name = "Aptos"
    body_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    body_style.font.size = Pt(10.5)


def add_title_block(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_after = Pt(2)
    run = p.add_run(TITLE)
    run.font.name = "Aptos Display"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(48, 76, 87)

    sub = document.add_paragraph()
    sub.space_after = Pt(14)
    sub_run = sub.add_run("供 Stitch 生成 UI 使用")
    sub_run.font.name = "Aptos"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sub_run.font.size = Pt(10.5)
    sub_run.font.color.rgb = RGBColor(110, 110, 110)


def add_summary_box(document):
    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    widths = [Inches(1.3), Inches(5.7)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    labels = ["产品名", "目标"]
    values = [
        "未至",
        "帮助用户在出发前，通过书籍和影视作品提前进入一座城市的故事与场景。",
    ]

    for i in range(2):
        label_cell = table.cell(i, 0)
        value_cell = table.cell(i, 1)
        set_cell_shading(label_cell, "E4ECE8")
        label_p = label_cell.paragraphs[0]
        label_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label_p.add_run(labels[i])
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.name = "Aptos"
        label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        value_p = value_cell.paragraphs[0]
        value_run = value_p.add_run(values[i])
        value_run.font.size = Pt(10.5)
        value_run.font.name = "Aptos"
        value_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_paragraph()


def add_section(document, heading, bullets):
    hp = document.add_paragraph()
    hp.space_before = Pt(8)
    hp.space_after = Pt(6)
    hr = hp.add_run(heading)
    hr.font.name = "Aptos"
    hr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    hr.font.size = Pt(13)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor(48, 76, 87)

    for bullet in bullets:
        p = document.add_paragraph(style="Body Accent")
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("• ")
        run.bold = True
        text_run = p.add_run(bullet)
        text_run.font.name = "Aptos"
        text_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        text_run.font.size = Pt(10.5)


def build_document():
    document = Document()
    set_page_margins(document.sections[0])
    configure_styles(document)
    add_title_block(document)
    add_summary_box(document)

    for heading, bullets in SECTIONS:
        add_section(document, heading, bullets)

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
